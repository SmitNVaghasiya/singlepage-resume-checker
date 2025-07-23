import os
import io
from typing import Optional, Tuple
import logging
from docx import Document
import PyPDF2
import pdfplumber
import filetype
from .config import settings

logger = logging.getLogger(__name__)

class FileProcessor:
    @staticmethod
    def detect_file_type(file_content: bytes, filename: str) -> str:
        try:
            kind = filetype.guess(file_content)
            if kind and kind.extension.lower() in settings.allowed_extensions_list:
                return kind.extension.lower()
            if filename:
                extension = filename.lower().split('.')[-1]
                if extension in settings.allowed_extensions_list:
                    return extension
            return 'unknown'
        except Exception as e:
            logger.error(f"Error detecting file type: {str(e)}")
            return 'unknown'
    
    @staticmethod
    def validate_file_size(file_content: bytes, filename: str) -> Tuple[bool, str]:
        file_size = len(file_content)
        max_size = settings.max_file_size
        if file_size > max_size:
            size_mb = file_size / (1024 * 1024)
            max_mb = max_size / (1024 * 1024)
            return False, f"File '{filename}' is too large ({size_mb:.1f}MB). Maximum allowed size is {max_mb}MB."
        return True, "File size is acceptable"
    
    @staticmethod
    def validate_pdf_pages(file_content: bytes, filename: str) -> Tuple[bool, str]:
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                page_count = len(pdf.pages)
                max_pages = settings.max_pdf_pages
                if page_count > max_pages:
                    return False, f"PDF '{filename}' has too many pages ({page_count}). Maximum allowed is {max_pages} pages."
                return True, f"PDF has {page_count} pages (within limit)"
        except Exception as e:
            logger.warning(f"Could not validate PDF pages for {filename}: {e}")
            return True, "Could not validate page count, proceeding"
    
    @staticmethod
    def validate_docx_pages(file_content: bytes, filename: str) -> Tuple[bool, str]:
        try:
            doc = Document(io.BytesIO(file_content))
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            estimated_pages = max(1, (paragraph_count * 5 + table_count * 100) // 500)
            max_pages = settings.max_docx_pages
            if estimated_pages > max_pages:
                return False, f"DOCX '{filename}' appears to have too many pages (estimated {estimated_pages}). Maximum allowed is {max_pages} pages."
            return True, f"DOCX estimated {estimated_pages} pages (within limit)"
        except Exception as e:
            logger.warning(f"Could not validate DOCX pages for {filename}: {e}")
            return True, "Could not validate page count, proceeding"
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> Tuple[str, bool]:
        text = ""
        success = False
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    success = True
                    logger.info("PDF text extracted successfully using pdfplumber")
                    return text.strip(), success
        except Exception as e:
            logger.warning(f"pdfplumber failed: {str(e)}")
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                success = True
                logger.info("PDF text extracted successfully using PyPDF2")
                return text.strip(), success
        except Exception as e:
            logger.error(f"PyPDF2 also failed: {str(e)}")
        try:
            decoded_text = file_content.decode('utf-8', errors='ignore')
            if len(decoded_text.strip()) > 50:
                success = True
                logger.info("PDF processed as plain text")
                return decoded_text.strip(), success
        except Exception as e:
            logger.error(f"Plain text extraction failed: {str(e)}")
        return "Error: Unable to extract text from PDF file", False
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> Tuple[str, bool]:
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            if text.strip():
                logger.info("DOCX text extracted successfully")
                return text.strip(), True
            else:
                return "Error: DOCX file appears to be empty", False
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            try:
                decoded_text = file_content.decode('utf-8', errors='ignore')
                if len(decoded_text.strip()) > 50:
                    logger.info("DOCX processed as plain text")
                    return decoded_text.strip(), True
            except Exception as e2:
                logger.error(f"Plain text extraction from DOCX failed: {str(e2)}")
            return f"Error: Unable to extract text from DOCX file - {str(e)}", False
    
    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> Tuple[str, bool]:
        encodings_to_try = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'ascii']
        for encoding in encodings_to_try:
            try:
                text = file_content.decode(encoding)
                if text.strip():
                    logger.info(f"TXT file decoded successfully using {encoding}")
                    return text.strip(), True
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error decoding TXT with {encoding}: {str(e)}")
                continue
        try:
            text = file_content.decode('utf-8', errors='ignore')
            if text.strip():
                logger.info("TXT file decoded with error handling")
                return text.strip(), True
        except Exception as e:
            logger.error(f"Final TXT decoding attempt failed: {str(e)}")
        return "Error: Unable to decode text file", False
    
    @staticmethod
    def process_file(file_content: bytes, filename: str) -> Tuple[str, bool, str]:
        if not file_content:
            return "Error: Empty file", False, "unknown"
        size_valid, size_msg = FileProcessor.validate_file_size(file_content, filename)
        if not size_valid:
            return size_msg, False, "unknown"
        file_type = FileProcessor.detect_file_type(file_content, filename)
        logger.info(f"Processing file: {filename}, detected type: {file_type}")
        if file_type == 'pdf':
            pages_valid, pages_msg = FileProcessor.validate_pdf_pages(file_content, filename)
            if not pages_valid:
                return pages_msg, False, file_type
            text, success = FileProcessor.extract_text_from_pdf(file_content)
            return text, success, file_type
        elif file_type == 'docx':
            pages_valid, pages_msg = FileProcessor.validate_docx_pages(file_content, filename)
            if not pages_valid:
                return pages_msg, False, file_type
            text, success = FileProcessor.extract_text_from_docx(file_content)
            return text, success, file_type
        elif file_type == 'txt' or file_type == 'unknown':
            text, success = FileProcessor.extract_text_from_txt(file_content)
            return text, success, file_type if file_type != 'unknown' else 'txt'
        else:
            return f"Error: Unsupported file type '{file_type}'. Allowed types: {', '.join(settings.allowed_extensions_list)}.", False, file_type
    
    @staticmethod
    def validate_content(text: str, content_type: str = "document") -> Tuple[bool, str]:
        if not text or not text.strip():
            return False, f"Error: {content_type} appears to be empty"
        if len(text.strip()) < 50:
            return False, f"Error: {content_type} is too short (minimum 50 characters required)"
        alphanumeric_chars = sum(c.isalnum() for c in text)
        if alphanumeric_chars < len(text) * 0.3:
            return False, f"Error: {content_type} contains too many non-readable characters"
        words = text.split()
        if content_type == "resume" and len(words) > settings.max_resume_words:
            return False, f"Error: {content_type} is too long (maximum {settings.max_resume_words} words allowed)"
        if content_type == "job description":
            if len(words) > settings.max_job_description_words:
                return False, f"Error: {content_type} is too long (maximum {settings.max_job_description_words} words allowed)"
            elif len(words) < settings.min_job_description_words:
                return False, f"Error: {content_type} is too short (minimum {settings.min_job_description_words} words required)"
        return True, "Valid content"
    
    @staticmethod
    def validate_job_description_content(text: str) -> Tuple[bool, str]:
        if not text or not text.strip():
            return False, "Error: Job description is empty"
        words = text.split()
        word_count = len(words)
        if word_count < settings.min_job_description_words:
            return False, f"Error: Job description is too short. Minimum {settings.min_job_description_words} words required, got {word_count}."
        if word_count > settings.max_job_description_words:
            return False, f"Error: Job description is too long. Maximum {settings.max_job_description_words} words allowed, got {word_count}."
        if word_count < 50:
            return False, "Error: Job description appears to be too short for meaningful analysis"
        alphanumeric_chars = sum(c.isalnum() for c in text)
        if alphanumeric_chars < len(text) * 0.3:
            return False, "Error: Job description contains too many non-readable characters"
        return True, f"Valid job description ({word_count} words)"

async def extract_text_from_file(file: "UploadFile") -> str:
    try:
        content = await file.read()
        text, success, file_type = FileProcessor.process_file(content, file.filename or "unknown")
        if not success:
            raise ValueError(text)
        return text
    except Exception as e:
        logger.error(f"File extraction failed: {str(e)}")
        raise ValueError(f"Failed to extract text from file: {str(e)}")